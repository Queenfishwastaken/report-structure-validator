from docx import Document
from io import BytesIO
import re


async def read_docx(file) -> list:
    try:
        file_content = await file.read()
        doc = Document(BytesIO(file_content))

        sections = []
        seen_sections = set()  # Для избежания дубликатов

        # Паттерны для определения заголовков
        HEADING_PATTERNS = [
            r'^\s*\d+[\.\)]\s+.+$',  # "1. Введение"
            r'^\s*[IVX]+[\.\)]\s+.+$',  # "I. Введение"
            r'^\s*\d+\.\d+\s+.+$',  # "1.1 Теория"
            r'^\s*[А-ЯЁ][А-ЯЁ\s]{3,}$',  # ВСЕ ЗАГЛАВНЫЕ буквы (4+ символа)
        ]

        # Ключевые слова для разделов
        SECTION_KEYWORDS = {
            'введение': ['введение', 'вступление', 'цель', 'задание', 'актуальность', 'постановка задачи'],
            'теория': ['теория', 'теоретическая', 'литературный обзор', 'теоретические основы'],
            'практика': ['практика', 'эксперимент', 'исследование', 'ход работы', 'методика', 'реализация'],
            'заключение': ['заключение', 'вывод', 'выводы', 'результаты', 'итоги'],
            'литература': ['литература', 'библиография', 'список', 'источники'],
            'содержание': ['оглавление', 'содержание', 'план'],
            'титульный': ['титульный', 'название', 'тема', 'лабораторная работа']
        }

        # Собираем все значимые заголовки
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text or len(text) < 2:
                continue

            # Очищаем текст от лишних символов
            clean_text = text.replace('**', '').replace('__', '').strip()

            # Определяем, может ли это быть заголовком
            is_heading = False
            section_type = None

            # 1. Проверяем стиль абзаца (самый надежный способ)
            style_name = paragraph.style.name.lower()
            if any(keyword in style_name for keyword in ['heading', 'title', 'header', 'заголовок']):
                is_heading = True

            # 2. Проверяем по паттернам
            if not is_heading:
                for pattern in HEADING_PATTERNS:
                    if re.match(pattern, clean_text, re.IGNORECASE):
                        is_heading = True
                        break

            # 3. Проверяем форматирование
            if not is_heading and paragraph.runs:
                # Проверяем жирный шрифт в первом run
                first_run = paragraph.runs[0]
                if first_run.bold and len(clean_text.split()) < 8:
                    is_heading = True

            # 4. Проверяем по ключевым словам
            if not is_heading:
                text_lower = clean_text.lower()
                for section_name, keywords in SECTION_KEYWORDS.items():
                    for keyword in keywords:
                        if keyword in text_lower:
                            is_heading = True
                            section_type = section_name
                            break
                    if is_heading:
                        break

            # Если это заголовок, добавляем его
            if is_heading:
                # Убираем номера и маркеры
                clean_heading = re.sub(r'^\s*[\dIVX\.\)]+\s*', '', clean_text)
                clean_heading = clean_heading.strip()

                if clean_heading and clean_heading not in seen_sections:
                    seen_sections.add(clean_heading)

                    # Для отладки можно добавить тип раздела
                    if section_type:
                        sections.append(f"{clean_heading}")  # Можно убрать комментарий для отладки
                    else:
                        sections.append(clean_heading)

        # Если нашли мало заголовков, добавляем первые параграфы, похожие на заголовки
        if len(sections) < 3:
            alternative_headings = []
            for paragraph in doc.paragraphs[:20]:  # Проверяем первые 20 параграфов
                text = paragraph.text.strip()
                if not text or len(text) < 10:
                    continue

                # Берем короткие параграфы (вероятно, заголовки)
                if len(text) < 150 and not text.endswith('.'):
                    words = text.split()
                    if len(words) < 10:  # Не более 10 слов
                        if text not in seen_sections:
                            alternative_headings.append(text)
                            seen_sections.add(text)

            sections.extend(alternative_headings[:5])

        # Если все еще мало, берем первые параграфы
        if not sections:
            sections = [p.text.strip() for p in doc.paragraphs if p.text.strip()][:10]

        # Очищаем от дубликатов и пустых строк
        final_sections = []
        for section in sections:
            if section and section.strip() and section not in final_sections:
                final_sections.append(section.strip())

        # Логируем для отладки
        print(f"Найдено разделов: {len(final_sections)}")
        for i, section in enumerate(final_sections, 1):
            print(f"  {i}. {section}")

        return final_sections if final_sections else ["Не удалось определить разделы"]

    except Exception as e:
        print(f"Ошибка в парсере: {str(e)}")
        raise Exception(f"Ошибка чтения файла: {str(e)}")